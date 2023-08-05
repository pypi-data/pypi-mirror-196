#!/usr/bin/env python
from docx import Document
from transformers import pipeline, AutoTokenizer, AutoModelForTokenClassification
import sys
from pathlib import Path
from os.path import join
from argparse import ArgumentParser
from . import _version
__version__ = _version.get_versions()['version']

class RemovePIIDocx(object):
    """
    Detect PII from a MS docx file.

    """
    model_named_entities = {"Jean-Baptiste/roberta-large-ner-english":{"PER": "Name",
                                                                       "LOC": "Location",
                                                                       "ORG": "Organization",
                                                                       "O"  : "Other",
                                                                       "MISC": "Miscellaneous",
                                                                       }
                            }
    model_entity_removal = {"Jean-Baptiste/roberta-large-ner-english":["PER"]}

    def __init__(self, docx_file=None, *args, **kwargs):
        self.docx_file = docx_file
        self.docx = Document(docx_file)
        self.pii_docx = Document()
        self._model = None
        self._model_name = None
        self._pii_complete = False

    @property
    def nlp_model(self):
        return self._model

    @nlp_model.setter
    def nlp_model(self, model_name):
        """this is actually not a model but a pipeline. But they don't need to know that :P

        TODO: should implement a check to make sure the model_name can perform NER.


        """
        self._model_name = model_name
        tokenizer=AutoTokenizer.from_pretrained(model_name)
        model = AutoModelForTokenClassification.from_pretrained(model_name)
        self._model = pipeline('ner', model=model, tokenizer=tokenizer, aggregation_strategy='simple')

    def pii(self):
        ent = self.model_named_entities[self._model_name]
        ent_count = {i:[] for i in ent.keys()}
        results = []
        texts = []
        # first loop through to collect the results of the PII detection.
        # after this we count the number of each to index the names so 'Peter Boyd' becomes 'Name1'
        # so-and-so becomes 'Name2', and so on...
        for i, para in enumerate(self.docx.paragraphs):
            res = self.nlp_model(para.text)
            
            results.append(res)
            texts.append(para.text)
            for r in res:
                ent_count[r['entity_group']].append(r['word'])
        ent_count = {i:list(set(ent_count[i])) for i in ent_count.keys()}
        replace_dic = {i:f"{ent[j]}{idx+1}" for j in ent_count.keys() for idx, i in enumerate(ent_count[j])}
        # second loop, go through each result and replace the text with PII removed.
        for res, text in zip(results, texts):
            for r in reversed(res):
                # only remove PII that are in desired categories (e.g. names only, but not location/org)
                if r['entity_group'] in self.model_entity_removal[self._model_name]:
                    text = text[:r['start']] + replace_dic[r['word']] + text[r['end']:]
            self.pii_docx.add_paragraph(text)
        self._pii_complete = True

    def write_pii_removed_file(self):
        if self._pii_complete:
            p = Path(self.docx_file)
            newp = join(p.parent, f"{p.stem}_PII_REM{p.suffix}")
            self.pii_docx.save(newp)
            print(f"File with PII removed has been saved to {newp}.")

        else:
            print("Have not performed PII removal yet!")

def cli():

    parser = ArgumentParser(description="***PII Removal for MS word documents***")
    parser.add_argument(dest='file_name', action='store',
                        help="Full path to the .docx file which will undergo PII removal.")
    parser.add_argument('--model-name', dest='model_name', action='store', default='Jean-Baptiste/roberta-large-ner-english', help="name of the (huggingface) model used to perform Named Entity Recognition (NER).\nDefault is 'Jean-Baptiste/roberta-large-ner-english'.")
    parser.add_argument('--version', '-v', action='version', version=f"%(prog)s version "+__version__)
    args = parser.parse_args()
    #TODO: assert that the docx file exists and is actually a docx file?
    piirem = RemovePIIDocx(args.file_name)
    piirem.nlp_model = args.model_name
    piirem.pii()
    piirem.write_pii_removed_file()

def main():
    print("Hello!")
    cli()

if __name__=="__main__":
    main()
